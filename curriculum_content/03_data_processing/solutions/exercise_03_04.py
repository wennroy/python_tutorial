# exercise_03_04.py - Reference Solution
# 练习 03-04: 办公自动化

import pandas as pd
from docx import Document
from docx.shared import Inches
import os

def generate_report():
    # 1. 准备数据
    data = {
        "Product": ["Apple", "Banana", "Orange", "Grape", "Mango"],
        "Sales": [5000, 3000, 4000, 2000, 6000]
    }
    df = pd.DataFrame(data)
    
    # 2. 计算统计信息
    total_sales = df["Sales"].sum()
    top_product = df.sort_values("Sales", ascending=False).iloc[0]
    
    print(f"Total Sales: {total_sales}")
    print(f"Top Product: {top_product['Product']}")
    
    # 3. 生成 Word 文档
    doc = Document()
    
    # 标题
    doc.add_heading('Weekly Sales Report', 0)
    
    # 摘要段落
    p = doc.add_paragraph('This week we achieved a total sales of ')
    p.add_run(f"${total_sales:,}").bold = True
    p.add_run('.')
    
    doc.add_paragraph(f"The best selling product is {top_product['Product']} with ${top_product['Sales']:,}.")
    
    # 表格标题
    doc.add_heading('Sales Details', level=1)
    
    # 创建表格 (行数 = 数据行数 + 1表头)
    table = doc.add_table(rows=len(df)+1, cols=2)
    table.style = 'Table Grid' # 添加边框样式
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Sales ($)'
    
    # 填充数据
    for i, row in df.iterrows():
        row_cells = table.rows[i+1].cells
        row_cells[0].text = str(row['Product'])
        row_cells[1].text = str(row['Sales'])
        
    # 保存
    filename = 'sales_report.docx'
    doc.save(filename)
    print(f"Report generated: {os.path.abspath(filename)}")

if __name__ == "__main__":
    # 注意：运行此脚本需要安装 python-docx
    # pip install python-docx
    try:
        generate_report()
    except ImportError:
        print("Error: 'python-docx' library not found. Please install it via 'pip install python-docx'")
